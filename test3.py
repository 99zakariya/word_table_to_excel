# -*- coding: utf-8 -*-
"""
Created on Sat Jul  4 02:16:18 2020

@author: zak
"""
import pandas as pd
from docx.api import Document
def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell
heads=["Type","Title","content","choice_1","choice_2","choice_3","choice_4","choice_5","correct_choice","Skill Name"]
data=[]

document = Document('C:/Users/mohammadzakar.meraj/Desktop/mcq Extract test/04_SQL_Basic_MCQ_Medium.docx')
for table in document.tables:
    sm=[]
    for row in table.rows:
        for cell in iter_unique_cells(row):
            for paragraph in cell.paragraphs:
                sm.append(paragraph.text)
    data.append(sm)

fullText = []
for para in document.paragraphs:
        fullText.append(para.text)
# print(fullText)
for p in fullText:
	if(p.isspace()):
		fullText.remove(p)
		#print('hello')
copy=fullText.copy()
for p in fullText:
	if not(p.startswith('Q')):
		copy.remove(p)
		#print('hello')
titil=[]
for s in range(len(copy)):
	titil.append((copy[s])[(1+(len(str(s+1)))):])
# print(titil)
    
mc=[]
ma=[]
topic=input("enter topic name") 
for k in data:
    name=""
    right=""
    new_data=[]
    if "MC" in k[1]:
        name="MCQ";
    else:
        name="MCA"
    new_data.append(name)
    new_data.append(titil.pop())
    new_data.append(k[0])
    ind=[15,19,23,27]
    ans="0"
    for p in range(4):
        new_data.append(k[ind[p]])
        
        if int(k[ind[p]+2]) >0:
            if ans is "0":
                ans = str(p+1)
            else:
                ans = ans+","+str(p+1)
    new_data.append("")
    new_data.append(ans)
    new_data.append(topic)
    if "MC" in k[1]:
        mc.append(new_data)
    else:
        ma.append(new_data)
    
    




mc_df = pd.DataFrame(mc,columns =heads) 
ma_df = pd.DataFrame(ma,columns =heads) 

writer = pd.ExcelWriter('C:/Users/mohammadzakar.meraj/Desktop/mcq Extract test/04_SQL_Basic_MCQ_Medium.xlsx', engine='xlsxwriter')

mc_df.to_excel(writer, sheet_name='Sheeta')

ma_df.to_excel(writer, sheet_name='Sheetb')



writer.save()
