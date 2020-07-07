# -*- coding: utf-8 -*-
"""
Created on Sat Jul  4 02:16:18 2020

@author: zak
"""

print('loding...')
import re
import time 
import os
import pandas as pd
from docx.api import Document
def unique_el(mc,dup):
    fc=[]
    for k in mc:
        flag=0
        for lol in fc:
            
            if k[2] == lol[2]:
                flag=1
        if flag ==0:
            fc.append(k)
        else:
            dup.append(k)
            
    return fc
def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell
heads=["Type","Title","content","choice_1","choice_2","choice_3","choice_4","choice_5","correct_choice","Skill Name"]

print("ensure the word file is in the same directry as python software!!")
path=(os.path.dirname(os.path.realpath(__file__)))
print(path)
while True:
    data=[]
    try:
        print("="*50,'press ','1 to extract data from word','2 to exit from this program ',"="*50,sep="\n")
        op=int(input(":"))
        if op ==1:
            
            filename=input("enter the name of the file it should be in same directory without extention")
            try:
                document = Document(path+'\\'+filename+'.docx')
                print("extracting from word file......")
                for table in document.tables:
                    sm=[]
                    for row in table.rows:
                        for cell in iter_unique_cells(row):
                            temp=cell.text
                            temp=temp.strip(' ')
                            sm.append(temp)
                            
                    data.append(sm)

                fullText = []
                for para in document.paragraphs:
                        fullText.append(para.text)
                
                r = re.compile('^Q([1-9]|\s+)')
                copy=[i for i in fullText if r.match(i)]
                titil=[]
                for s in range(len(copy)):
                    titil.append((copy[s])[(1+(len(str(s+1)))):])
                print("extracting done")
                mc=[]
                ma=[]
                err=[]
                dup=[]
                title=titil.reverse()
                topic=input("enter topic name:")
                print("formating extracted data....")
                con=0;
                for k in data:
                    con=con+1
                    try:
                        ind=[]
                        if k.count('Grade')>0:
                           s =k.index('Grade')+2
                           temp=[s+(p*4) for p in list(range(4))]
                           ind.extend(temp)
                        else:
                            raise Exception("Grade not found ")
                        name=""
                        right=""
                        new_data=[]
                        if "MC" in k[1]:
                            name="MCQ";
                        else:
                            name="MCA"
                        new_data.append(name)
                        if titil:
                            new_data.append(titil.pop())
                        else:
                            new_data.append('')
                        new_data.append(k[0])
                        
                        ans="0"
                        
                        for p in range(4):
                            new_data.append(k[ind[p]])
                            #print("hello",ind,"hello")
                            if float(k[ind[p]+2]) >0:
                                if ans == "0":
                                    ans = str(p+1)
                                else:
                                    ans = ans+","+str(p+1)
                        new_data.append("")
                        new_data.append(ans)
                        new_data.append(topic)
                        if "MC" in k[1]:
                            mc.append(new_data)
                        else:
                            ma.append(new_data)
                    except Exception as ex:
                        print("😯      enconeuntered problem in question no ",con,'  : ',ex)
                        err.append(k)
                print("formating done!")
                print("removing dublicate.....")      
               
                fmc=unique_el(mc,dup)
                fma=unique_el(ma,dup)
                print("done removing dublicate!")
                print("wrinting to excel file")

                mc_df = pd.DataFrame(fmc,columns =heads) 
                ma_df = pd.DataFrame(fma,columns =heads)
                dup_df = pd.DataFrame(dup,columns =heads)
                err_df=pd.DataFrame(err)

                writer = pd.ExcelWriter(path+'\\'+filename+'.xlsx', engine='xlsxwriter')

                mc_df.to_excel(writer, sheet_name='Multiple choice question',index=False)

                ma_df.to_excel(writer, sheet_name='Multiple choice answer',index=False)
                dup_df.to_excel(writer, sheet_name='duplicate questions',index=False)
                err_df.to_excel(writer, sheet_name='Question having error',index=False)



                writer.save()
                print("writing done!")
                print("😇 task completed ",len(data),' question found  - saved you ',2*len(data),' min')
                time.sleep(1)
            except Exception as ex:
                print(' somthing went wrong 😟  ',ex)
        elif op ==2:
            print("bye ... bye .. 🖐️")
            time.sleep(2)
            #exit(0)
            break
        else:
            raise Exception()
    except :
        print("please chose correct option 😒  ")
