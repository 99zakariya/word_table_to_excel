
import tkinter as tk 
import tkinter.scrolledtext as st
from tkinter import filedialog
from threading import Thread
from time import sleep
import os
import re
import time 
import pandas as pd
from docx.api import Document
from tkinter import simpledialog
dira=False
close=False
def selfile_fn():
    global dira
    sel.delete(0, last=len( sel.get()))
    sav.delete(0, last=len( sav.get()))
    filename = filedialog.askopenfilename(initialdir = "/",title = "Select word file ",filetypes = (("word files","*.docx"),("all files","*.*")))
    sel.insert( 0, filename )
    x=filename.rindex('.')
    sav.insert(0,filename[:x]+'.xlsx')
    dira=False
def selfol_fn():
    global dira
    sel.delete(0, last=len( sel.get()))
    sav.delete(0, last=len( sav.get()))
    filename = filedialog.askdirectory(initialdir = "/",title = "Select folder")
    sel.insert( 0, filename )
    sav.insert( 0, filename )
    dira=True
def savfol_fn():
    global dira
    sav.delete(0, last=len( sav.get()))
    filename="*"
    print(dira)
    if dira:
        filename = filedialog.askdirectory(initialdir = "/",title = "Select folder to save in")
    else:
        filename = filedialog.asksaveasfilename(initialdir = "/",title = "where to save file",defaultextension='.xlsx')
    sav.insert( 0, filename )
    
def con_fn():
    if len(sel.get())>0 and len(sav.get())>0:
        global dira,file_list
        
        onlyfiles=[];
        path=sel.get()
        spath=sav.get()
        if  dira:
            
            onlyfiles = [f for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]
            onlyfiles = [f for f in onlyfiles if f.endswith(".docx")]
            for k in range(len(onlyfiles)):
                onlyfiles[k]=path+"/"+onlyfiles[k]
                file_list.append((onlyfiles[k],spath,dira))
                
            
        else:
            file_list.append((path,spath,dira))
            
       # print(file_list,":",len(file_list))
        
    else:
        
        print("fail")
def on_closing():
    global close,win
    close=True
    win.destroy()
def mes_stu(mas):
     staus_area['state']=tk.NORMAL
     staus_area.insert(tk.INSERT,"\n"+mas)
     staus_area['state']=tk.DISABLED
     staus_area.see("end")
def mes_fil(mas):
     files_area['state']=tk.NORMAL
     files_area.insert(tk.INSERT,"\n"+mas)
     files_area['state']=tk.DISABLED
     files_area.see("end")
def unique_el(mc,dup):
    fc=[]
    for k in mc:
        flag=0
        for lol in fc:
            
            if k[2] == lol[2]:
                flag=1
        if flag ==0:
            fc.append(k)
        else:
            dup.append(k)
            
    return fc
def iter_unique_cells(row):
    """Generate cells in *row* skipping empty grid cells."""
    prior_tc = None
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is prior_tc:
            continue
        prior_tc = this_tc
        yield cell

def dil(msg):
    newWin = tk.Tk()

    #But make it invisible

    newWin.withdraw()

    #Now this works without throwing an exception:

    retVal = simpledialog.askstring("Enter Value",msg,parent=newWin)

    #Destroy the temporary "parent"

    newWin.destroy()
    return retVal
    

def mainlogic():
    global close
    fn=""
    while True:
        if close:
            break
        data=[]
        try:
            if file_list:
                
                #headst=["Type","Title","content","choice_1","choice_2","choice_3","choice_4","choice_5","correct_choice","Skill Name"]
                lent=0
                heads=["Type","Title","content"]
                sel=file_list.pop()
                fn=sel[0][sel[0].rindex('/'):]
                sadir=sel[1]
                if sel[2]:
                    sadir=sadir+fn[:fn.rindex('.')]+".xlsx"
                sn=sadir[sadir.index('/'):]
                
                mes_fil("🤓  "+fn+" starting...")
                try:
                    document = Document(sel[0])
                    mes_stu("👉  extracting from  file......")
                    for table in document.tables:
                        sm=[]
                        for row in table.rows:
                            for cell in iter_unique_cells(row):
                                temp=cell.text
                                temp=temp.replace('\xa0', '')
                                temp=str(temp).strip(' ')
                                
                                sm.append(temp)
                                
                        data.append(sm)

                    fullText = []
                    for para in document.paragraphs:
                            fullText.append(para.text)
                    
                    r = re.compile('^Q([1-9]|\s+)')
                    copy=[i for i in fullText if r.match(i)]
                    titil=[]
                    for s in range(len(copy)):
                        titil.append((copy[s])[(1+(len(str(s+1)))):])
                    mes_stu("✔ extracting done  !")
                    mc=[]
                    ma=[]
                    err=[]
                    dup=[]
                    title=titil.reverse()
                    topic = dil( "enter topic name of"+fn+":")
                    #topic=input("enter topic name:")
                    mes_stu("👉  formating extracted data of ....")
                    con=0;
                   
                        #print(len(data[l]),data[l])
                    data_cor=[]
                    for k in data:
                        con=con+1
                        lent_t=0
                        
                        try:
                            if k.count('Grade')>0:
                                s =k.index('Grade')+2
                                try:
                                    an=s+2
                                    while True:
                                        
                                        float(k[an])
                                        an=an+4
                                        lent_t=lent_t+1
                                except:
                                    temp=[s+(p*4) for p in list(range(lent_t))]
                                k.append(con)
                                k.append(s)
                                k.append(temp)
                                k.append(str(lent_t))
                                if lent_t>lent:
                                    lent=lent_t
                            else:
                                raise Exception("Grade not found")
                        except Exception as ex:
                            mes_stu("😒 enconeuntered problem in  quuestion no "+str(con)+'  : '+str(ex))
                            err.append(k)
                        else:
                            data_cor.append(k)
                    for k in range(lent):
                        #print("here",con)
                        heads.append("choice_"+str(k+1))
                    heads.extend(["correct_choice","Skill Name"])
                            
                    for k in data_cor:
                        
                        try:
                           
                            ind=k[-2]
                            #print(k.count('Grade'),'count')
                            
                            name=""
                            right=""
                            new_data=[]
                            if "MC" in k[1]:
                                name="MCQ";
                            else:
                                name="MCA"
                            new_data.append(name)
                            if titil:
                                new_data.append(titil.pop())
                            else:
                                new_data.append('')
                            new_data.append(k[0])
                            
                            ans="0"
                            
                            for p in range(len(ind)):
                                new_data.append(k[ind[p]])
                                #print("hello",ind,"hello")
                                if float(k[ind[p]+2]) >0:
                                    if ans == "0":
                                        ans = str(p+1)
                                    else:
                                        ans = ans+","+str(p+1)
                            for p in range(len(ind),lent):
                                new_data.append("")
                            new_data.append(ans)
                            new_data.append(topic)
                            if ans == "0":
                                raise Exception(" none of the ans are correct")
                            elif "MCQ" == new_data[0] and len(ans)>1 :
                                raise Exception(" MCQ type is having multiple ans correct")
                            elif "MCA" == new_data[0] and len(ans)==1:
                                raise Exception(" MCA type is have only one ans correct")   
                            if "MC" in k[1]:
                                mc.append(new_data)
                            else:
                                ma.append(new_data)
                        except Exception as ex:
                            mes_stu("😒 enconeuntered problem in  quuestion no "+str(k[-4])+'  : '+str(ex))
                            err.append(k)
                    mes_stu("✔ formating done  !")
                    mes_stu("👉  removing dublicate .....")      
                   
                    fmc=unique_el(mc,dup)
                    fma=unique_el(ma,dup)
                    mes_stu("✔ done removing dublicate "+str(len(dup))+"!")
                    mes_stu("👉   wrinting to  file...")

                    mc_df = pd.DataFrame(fmc,columns =heads) 
                    ma_df = pd.DataFrame(fma,columns =heads)
                    dup_df = pd.DataFrame(dup,columns =heads)
                    err_df=pd.DataFrame(err)

                    writer = pd.ExcelWriter(sadir, engine='xlsxwriter')

                    mc_df.to_excel(writer, sheet_name='Multiple choice question',index=False)

                    ma_df.to_excel(writer, sheet_name='Multiple choice answer',index=False)
                    dup_df.to_excel(writer, sheet_name='duplicate questions',index=False)
                    err_df.to_excel(writer, sheet_name='Question having error',index=False)



                    writer.save()
                    mes_stu("✔ writing done!")
                      
                    mes_stu("😊 task completed "+str(len(data))+' question found  - saved you '+str(2*len(data))+' minutes !   \n' )
                    
                    
                except Exception as ex:
                    mes_stu('😒 somthing went wrong '+str(ex))
                    mes_fil("😲  "+fn+" was not done!")
                
                    
                else:    
                    mes_fil("😇  "+fn+" done!")
                  
                
                    
            
        except Exception as ex:
            mes_stu('😒 somthing went wrong end '+str(ex))
            mes_fil("😲  "+fn+" was not done!")
        
        
    print("out")       

    
       
file_list=[]

# Creating tkinter window 
win = tk.Tk() 
win.title("word_table_to_excel V 3.1") 
  
# Title Label 
tk.Label(win,  
         text = "File or folder name",  
         font = ("Times New Roman", 15),  
         background = '#EDF568',  
         foreground = "black").grid(column = 0, 
                                    row = 0)
tk.Label(win,  
         text = "Where to save the file",  
         font = ("Times New Roman", 15),  
         background = '#EDF568',  
         foreground = "black").grid(column = 0, 
                                    row = 3)

sel = tk.Entry(win,font = ("Times New Roman", 13) ,width=30) 
sav = tk.Entry(win,font = ("Times New Roman", 13),width=30)

sel_file_in = tk.Button(win, text='Select File ',font = ("Times New Roman", 15),background = '#B5DF4D', width=13,  padx=2 , pady=2,command=selfile_fn).grid(row=0,  pady = 30,column=3)
sel_fol_in = tk.Button(win, text='Select Folder ', font = ("Times New Roman", 15),background = '#5FEAC2',width=13,  padx=2 ,pady=2,command=selfol_fn).grid(row=0, pady = 30, column=4)
sel_fol_out = tk.Button(win, text='Select Folder ', font = ("Times New Roman", 15),background = '#5FEAC2',width=13,  padx=2 ,pady=2,command=savfol_fn).grid(row=3, column=3)
con_start = tk.Button(win, text='Start Conversion ', font = ("Times New Roman", 15),background = '#5BF54C',width=15,  padx=2 ,pady=2,command=con_fn).grid(row=4, pady = 30, column=1)
sel.grid(row=0, column=1) 
sav.grid(row=3, column=1)
  
# Creating scrolled text area 
# widget with Read only by 
# disabling the state
tk.Label(win,  
         text = "Files worked on",  
         font = ("Times New Roman", 15),  
         background = '#EDF568',  
         foreground = "black").grid(column = 0,columnspan=2 )
files_area = st.ScrolledText(win, 
                            width = 45,  
                            height = 15,  
                            font = ("Times New Roman", 
                                    15)) 
  
files_area.grid(column = 0,columnspan=2, pady = 10, padx = 10) 
  
# Inserting Text which is read only 
files_area.insert(tk.INSERT, 
"""\ """)

  
# Making the text read only 
files_area.configure(state ='disabled') 
tk.Label(win,  
         text = "Status",  
         font = ("Times New Roman", 15),  
         background = '#EDF568',  
         foreground = "black").grid(column = 3,row=5,columnspan=2 )
staus_area = st.ScrolledText(win, 
                            width = 45,  
                            height = 15,  
                            font = ("Times New Roman", 
                                    15)) 
  
staus_area.grid(column = 3,row=6,columnspan=2, pady = 10, padx = 10) 
  
# Inserting Text which is read only 
staus_area.insert(tk.INSERT, 
"""\ """)

  
# Making the text read only 
staus_area.configure(state ='disabled')
thread = Thread(target = mainlogic)
thread.start()
win.protocol("WM_DELETE_WINDOW", on_closing)
win.mainloop() 
thread.join()
